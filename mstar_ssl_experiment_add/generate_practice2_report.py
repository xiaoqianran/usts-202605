#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Practice 2 report generator for MSTAR semi-supervised learning experiment.

Put this file in the project root, then run:

    python generate_practice2_report.py --runs runs --out 组号-人工智能实训II-实践2.docx

Expected directory structure:

runs/
  01_supervised_full/
    metrics.json
    history.csv
    confusion_matrix.csv
  02_supervised_10percent/
    metrics.json
    history.csv
    confusion_matrix.csv
  03_fixmatch_10percent/
    metrics.json
    history.csv
    confusion_matrix.csv

The script will generate:
  - a compact Word report for Practice 2
  - accuracy/recovery plot
  - loss/accuracy training curves
  - confusion matrices
  - FixMatch flowchart
"""
from __future__ import annotations

import argparse
import json
import math
import os
import sys
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# -----------------------------
# Basic configuration
# -----------------------------

CLASS_NAMES_DEFAULT = [
    "2S1", "BMP2", "BRDM_2", "BTR60", "BTR70",
    "D7", "T62", "T72", "ZIL131", "ZSU_23_4"
]

DATASET_TRAIN_COUNTS = {
    "2S1": 299, "BMP2": 233, "BRDM_2": 298, "BTR60": 256, "BTR70": 233,
    "D7": 299, "T62": 298, "T72": 232, "ZIL131": 299, "ZSU_23_4": 299,
}

DATASET_TEST_COUNTS = {
    "2S1": 274, "BMP2": 195, "BRDM_2": 274, "BTR60": 195, "BTR70": 196,
    "D7": 274, "T62": 273, "T72": 196, "ZIL131": 274, "ZSU_23_4": 274,
}

DISPLAY_NAMES = {
    "full": "全标签监督",
    "limited": "10%标签监督",
    "fixmatch": "FixMatch半监督",
}

SHORT_DISPLAY_NAMES = {
    "full": "Full",
    "limited": "10% Sup.",
    "fixmatch": "FixMatch",
}

PLOT_NAMES = {
    "full": "Full supervised",
    "limited": "10% supervised",
    "fixmatch": "10% FixMatch",
}

ORDER = ["full", "limited", "fixmatch"]

# -----------------------------
# Font helpers
# -----------------------------

def setup_matplotlib_font() -> None:
    """Try to use a Chinese-capable font if available."""
    candidates = [
        "Microsoft YaHei", "SimHei", "SimSun", "Noto Sans CJK SC",
        "Noto Serif CJK SC", "WenQuanYi Micro Hei", "Arial Unicode MS"
    ]
    available = {f.name for f in font_manager.fontManager.ttflist}
    for name in candidates:
        if name in available:
            plt.rcParams["font.sans-serif"] = [name]
            plt.rcParams["font.family"] = "sans-serif"
            break
    plt.rcParams["axes.unicode_minus"] = False


def set_run_font(run, size: float = 10.5, bold: bool = False, name: str = "SimSun") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text, size: float = 8.2, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str = "EDEDED") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

# -----------------------------
# Reading experiment outputs
# -----------------------------

def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def find_experiment_dirs(runs_dir: Path) -> Dict[str, Path]:
    """Discover full supervised / limited supervised / FixMatch folders."""
    if not runs_dir.exists():
        raise FileNotFoundError(f"runs directory not found: {runs_dir}")

    found: Dict[str, Path] = {}
    for d in sorted([p for p in runs_dir.iterdir() if p.is_dir()]):
        metrics_path = d / "metrics.json"
        if not metrics_path.exists():
            continue
        try:
            m = load_json(metrics_path)
        except Exception:
            continue
        method = str(m.get("method", "")).lower()
        ratio = float(m.get("label_ratio", -1))
        name = d.name.lower()
        if method == "fixmatch" or "fixmatch" in name:
            found["fixmatch"] = d
        elif method == "supervised" and (abs(ratio - 1.0) < 1e-8 or "full" in name):
            found["full"] = d
        elif method == "supervised" and ratio < 1.0:
            found["limited"] = d

    missing = [k for k in ORDER if k not in found]
    if missing:
        raise RuntimeError(
            "Cannot find required experiment folders in runs/. Missing: " + ", ".join(missing) + "\n"
            "Please make sure metrics.json exists under each experiment folder."
        )
    return found


def read_experiment(d: Path) -> dict:
    data = {
        "dir": d,
        "metrics": load_json(d / "metrics.json"),
        "history": None,
        "confusion": None,
        "classes": CLASS_NAMES_DEFAULT,
    }
    history_path = d / "history.csv"
    if history_path.exists():
        data["history"] = pd.read_csv(history_path)
    cm_path = d / "confusion_matrix.csv"
    if cm_path.exists():
        cm_df = pd.read_csv(cm_path)
        # Expected first column: true/pred, remaining columns = predicted classes.
        if cm_df.shape[1] >= 2:
            classes = list(cm_df.columns[1:])
            data["classes"] = classes
            mat = cm_df.iloc[:, 1:].to_numpy(dtype=float)
            data["confusion"] = mat
    return data


def pct(x: float) -> float:
    return float(x) * 100.0

# -----------------------------
# Figure generation
# -----------------------------

def make_accuracy_recovery_figure(exps: Dict[str, dict], fig_dir: Path) -> Path:
    fig_dir.mkdir(parents=True, exist_ok=True)
    vals = [pct(exps[k]["metrics"].get("best_test_acc", 0)) for k in ORDER]
    labels = [PLOT_NAMES[k] for k in ORDER]
    full, limited, fix = vals
    drop = full - limited
    gain = fix - limited
    recover = gain / drop * 100 if drop > 0 else 0

    fig, axes = plt.subplots(1, 2, figsize=(9.8, 3.2), dpi=220)

    ax = axes[0]
    bars = ax.bar(labels, vals)
    ax.set_ylim(0, 105)
    ax.set_ylabel("Test accuracy (%)")
    ax.set_title("Accuracy comparison")
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width()/2, v + 1.2, f"{v:.2f}%", ha="center", fontsize=8)
    ax.tick_params(axis="x", labelrotation=15)

    ax = axes[1]
    ax.plot([0, 1, 2], vals, marker="o")
    ax.set_xticks([0, 1, 2])
    ax.set_xticklabels(["Full", "10% Sup.", "FixMatch"], rotation=15)
    ax.set_ylim(0, 105)
    ax.set_ylabel("Test accuracy (%)")
    ax.set_title("Accuracy recovery")
    ax.grid(True, linestyle="--", alpha=0.35)
    ax.annotate(f"Drop {drop:.2f}pp", xy=(0.5, (full+limited)/2), xytext=(0.15, (full+limited)/2-18),
                arrowprops=dict(arrowstyle="->", lw=1), fontsize=8)
    ax.annotate(f"Recover {gain:.2f}pp\n{recover:.1f}% of loss", xy=(1.55, (limited+fix)/2), xytext=(1.25, (limited+fix)/2+12),
                arrowprops=dict(arrowstyle="->", lw=1), fontsize=8)
    for i, v in enumerate(vals):
        ax.text(i, v + 1.5, f"{v:.2f}%", ha="center", fontsize=8)

    fig.tight_layout()
    out = fig_dir / "fig2_accuracy_recovery.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return out


def get_col(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    for c in candidates:
        if c in df.columns:
            return c
    return None


def make_training_curves_figure(exps: Dict[str, dict], fig_dir: Path) -> Path:
    """Create one compact 2x2 figure with train/test loss and train/test accuracy."""
    fig, axes = plt.subplots(2, 2, figsize=(9.8, 5.3), dpi=220)
    curve_items = [
        (axes[0, 0], "train_loss", ["train_loss", "loss", "loss_x"], "Train loss"),
        (axes[0, 1], "test_loss", ["test_loss", "val_loss"], "Test loss"),
        (axes[1, 0], "train_acc", ["train_acc", "labeled_train_acc"], "Train accuracy"),
        (axes[1, 1], "test_acc", ["test_acc", "val_acc"], "Test accuracy"),
    ]
    for ax, _key, candidates, title in curve_items:
        plotted = False
        for k in ORDER:
            df = exps[k].get("history")
            if df is None or df.empty:
                continue
            epoch_col = "epoch" if "epoch" in df.columns else df.columns[0]
            val_col = get_col(df, candidates)
            if val_col is None:
                continue
            y = df[val_col].to_numpy(dtype=float)
            if "acc" in val_col:
                y = y * 100
            ax.plot(df[epoch_col], y, label=PLOT_NAMES[k], linewidth=1.3)
            plotted = True
        ax.set_title(title)
        ax.set_xlabel("epoch")
        ax.grid(True, linestyle="--", alpha=0.35)
        if "accuracy" in title.lower():
            ax.set_ylabel("%")
            ax.set_ylim(0, 105)
        if plotted:
            ax.legend(fontsize=7)
        else:
            ax.text(0.5, 0.5, "history.csv not found", ha="center", va="center", transform=ax.transAxes)
    fig.tight_layout()
    out = fig_dir / "fig2_training_curves.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return out


def make_individual_curve_figures(exps: Dict[str, dict], fig_dir: Path) -> List[Path]:
    """Optional individual curves for each experiment, saved for appendix/backup."""
    outs = []
    for k in ORDER:
        df = exps[k].get("history")
        if df is None or df.empty:
            continue
        fig, axes = plt.subplots(1, 2, figsize=(8.0, 2.9), dpi=180)
        epoch_col = "epoch" if "epoch" in df.columns else df.columns[0]
        # loss
        train_loss_col = get_col(df, ["train_loss", "loss", "loss_x"])
        test_loss_col = get_col(df, ["test_loss", "val_loss"])
        if train_loss_col:
            axes[0].plot(df[epoch_col], df[train_loss_col], label="train loss")
        if test_loss_col:
            axes[0].plot(df[epoch_col], df[test_loss_col], label="test loss")
        axes[0].set_title(f"{PLOT_NAMES[k]}: Loss")
        axes[0].set_xlabel("epoch")
        axes[0].grid(True, linestyle="--", alpha=0.35)
        axes[0].legend(fontsize=7)
        # acc
        train_acc_col = get_col(df, ["train_acc", "labeled_train_acc"])
        test_acc_col = get_col(df, ["test_acc", "val_acc"])
        if train_acc_col:
            axes[1].plot(df[epoch_col], df[train_acc_col] * 100, label="train acc")
        if test_acc_col:
            axes[1].plot(df[epoch_col], df[test_acc_col] * 100, label="test acc")
        axes[1].set_title(f"{PLOT_NAMES[k]}: Accuracy")
        axes[1].set_xlabel("epoch")
        axes[1].set_ylabel("%")
        axes[1].set_ylim(0, 105)
        axes[1].grid(True, linestyle="--", alpha=0.35)
        axes[1].legend(fontsize=7)
        fig.tight_layout()
        out = fig_dir / f"curve_{k}.png"
        fig.savefig(out, bbox_inches="tight")
        plt.close(fig)
        outs.append(out)
    return outs


def normalize_cm(cm: np.ndarray) -> np.ndarray:
    row_sum = cm.sum(axis=1, keepdims=True)
    row_sum[row_sum == 0] = 1
    return cm / row_sum


def make_confusion_matrices_figure(exps: Dict[str, dict], fig_dir: Path) -> Path:
    fig, axes = plt.subplots(1, 3, figsize=(12.0, 3.8), dpi=220)
    for ax, k in zip(axes, ORDER):
        cm = exps[k].get("confusion")
        classes = exps[k].get("classes") or CLASS_NAMES_DEFAULT
        if cm is None:
            ax.axis("off")
            ax.text(0.5, 0.5, "confusion_matrix.csv\nnot found", ha="center", va="center", transform=ax.transAxes)
            ax.set_title(PLOT_NAMES[k])
            continue
        cm_norm = normalize_cm(cm)
        im = ax.imshow(cm_norm, vmin=0, vmax=1)
        ax.set_title(PLOT_NAMES[k], fontsize=10)
        ax.set_xticks(range(len(classes)))
        ax.set_yticks(range(len(classes)))
        ax.set_xticklabels(classes, rotation=45, ha="right", fontsize=6)
        ax.set_yticklabels(classes, fontsize=6)
        ax.set_xlabel("Predicted class", fontsize=8)
        ax.set_ylabel("True class", fontsize=8)
        for i in range(cm_norm.shape[0]):
            for j in range(cm_norm.shape[1]):
                val = cm_norm[i, j]
                if val > 0.65 or (i == j and val > 0.35):
                    color = "white"
                else:
                    color = "black"
                if i == j or val >= 0.10:
                    ax.text(j, i, f"{val:.2f}", ha="center", va="center", fontsize=5.4, color=color)
    cbar = fig.colorbar(im if 'im' in locals() else axes[-1].imshow(np.zeros((2,2))), ax=axes.ravel().tolist(), shrink=0.78)
    cbar.set_label("Row-normalized proportion", fontsize=8)
    out = fig_dir / "fig2_confusion_matrices.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return out


def make_fixmatch_flowchart(fig_dir: Path) -> Path:
    fig, ax = plt.subplots(figsize=(10.0, 4.1), dpi=220)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")

    def box(x, y, w, h, text, fc="#F5F7FA"):
        patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.04,rounding_size=0.08",
                               linewidth=1.1, edgecolor="#333333", facecolor=fc)
        ax.add_patch(patch)
        ax.text(x + w/2, y + h/2, text, ha="center", va="center", fontsize=9)
        return patch

    def arrow(x1, y1, x2, y2):
        arr = FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="->", mutation_scale=12, linewidth=1.0, color="#333333")
        ax.add_patch(arr)

    # Supervised branch
    box(0.4, 3.55, 1.4, 0.55, "Labeled image\n(x, y)", "#EAF4FF")
    box(2.2, 3.55, 1.4, 0.55, "Weak aug.", "#EAF4FF")
    box(4.0, 3.55, 1.4, 0.55, "Model", "#EAF4FF")
    box(5.8, 3.55, 1.7, 0.55, "CE loss\nLx", "#EAF4FF")
    arrow(1.8, 3.825, 2.2, 3.825)
    arrow(3.6, 3.825, 4.0, 3.825)
    arrow(5.4, 3.825, 5.8, 3.825)

    # Unlabeled branch
    box(0.4, 1.55, 1.4, 0.55, "Unlabeled image\nu", "#FFF3E6")
    box(2.2, 2.25, 1.4, 0.55, "Weak aug.", "#FFF3E6")
    box(4.0, 2.25, 1.4, 0.55, "Model", "#FFF3E6")
    box(5.8, 2.25, 1.7, 0.55, "High-conf.\npseudo-label q", "#FFF3E6")
    box(2.2, 0.85, 1.4, 0.55, "Strong aug.", "#FFF3E6")
    box(4.0, 0.85, 1.4, 0.55, "Model", "#FFF3E6")
    box(5.8, 0.85, 1.7, 0.55, "Consistency\nloss Lu", "#FFF3E6")

    arrow(1.8, 1.825, 2.2, 2.525)
    arrow(3.6, 2.525, 4.0, 2.525)
    arrow(5.4, 2.525, 5.8, 2.525)
    arrow(1.8, 1.825, 2.2, 1.125)
    arrow(3.6, 1.125, 4.0, 1.125)
    arrow(5.4, 1.125, 5.8, 1.125)
    arrow(6.65, 2.25, 6.65, 1.40)

    box(8.0, 2.15, 1.65, 0.72, "Total loss\nL = Lx + λuLu", "#E9F8EF")
    arrow(7.5, 3.825, 8.25, 2.87)
    arrow(7.5, 1.125, 8.25, 2.15)

    ax.text(0.4, 4.55, "FixMatch pipeline: labeled CE loss + high-confidence pseudo-label + consistency regularization", fontsize=11, weight="bold")
    ax.text(5.85, 1.95, "threshold τ=0.95", fontsize=8)
    ax.text(8.05, 1.82, "λu=1.0", fontsize=8)

    fig.tight_layout()
    out = fig_dir / "fig2_fixmatch_flowchart.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return out

# -----------------------------
# DOCX generation
# -----------------------------

def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=16, bold=True, name="SimHei")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level > 1 else 8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=12 if level == 1 else 10.5, bold=True, name="SimHei" if level == 1 else "SimSun")


def add_para(doc: Document, text: str, indent: bool = True, size: float = 9.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    r = p.add_run(text)
    set_run_font(r, size=size, name="SimSun")


def add_table(doc: Document, caption: str, headers: List[str], rows: List[List], font_size: float = 7.8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(caption)
    set_run_font(r, size=9, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, size=font_size, bold=True)
        shade_cell(table.rows[0].cells[j])
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, size=font_size)


def add_image(doc: Document, path: Path, caption: str, width: float = 6.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(caption)
    set_run_font(r, size=8.5)


def add_footer_page_number(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build_report(exps: Dict[str, dict], figs: Dict[str, Path], out_docx: Path, group_name: str = "组号") -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("《人工智能开发实训 II》非卷面试题答题文档")
    set_run_font(r, size=8.5, bold=True)
    add_footer_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(0)

    # Metrics
    acc_full = pct(exps["full"]["metrics"]["best_test_acc"])
    acc_limited = pct(exps["limited"]["metrics"]["best_test_acc"])
    acc_fix = pct(exps["fixmatch"]["metrics"]["best_test_acc"])
    drop = acc_full - acc_limited
    gain = acc_fix - acc_limited
    gap = acc_full - acc_fix
    recover = gain / drop * 100 if drop > 0 else 0.0
    labeled = int(exps["limited"]["metrics"].get("num_labeled", 268))
    unlabeled = int(exps["fixmatch"]["metrics"].get("num_unlabeled", 2478))
    total_train = int(exps["full"]["metrics"].get("num_labeled", 2746))
    label_ratio = labeled / total_train * 100 if total_train else 0

    add_title(doc, "实践二  深度半监督学习实践")

    add_heading(doc, "2.1 实验内容与数据集", 1)
    add_para(doc, f"本实践基于MSTAR灰度SAR目标图像完成10类分类实验，目的是观察标签有限对深度分类网络性能的影响，并采用FixMatch半监督学习方法利用未标注样本缓解性能下降。训练集共{total_train}张，测试集共2425张；有限标签设置下仅保留{labeled}张有标签样本，占训练集{label_ratio:.2f}%，满足少于原训练样本1/10的要求。")

    add_table(
        doc,
        "表2-1 数据划分与实验设置",
        ["项目", "样本数", "说明"],
        [
            ["训练集", total_train, "10类MSTAR SAR图像"],
            ["测试集", 2425, "用于最终分类精度评价"],
            ["有限标签集", labeled, f"约10%标签，实际占比{label_ratio:.2f}%"],
            ["无标签集", unlabeled, "FixMatch半监督阶段使用"],
        ],
        font_size=7.8,
    )

    add_heading(doc, "2.2 基线与有限标签实验", 1)
    add_para(doc, "首先使用全部训练标签训练SmallResNet作为监督学习基线；随后仅使用268张有标签样本重新训练同一网络，其余训练样本不参与监督训练。两组实验用于验证标签数量减少对深度模型泛化性能的影响。")
    add_table(
        doc,
        "表2-2 基线与有限标签结果对比",
        ["方法", "有标签样本", "无标签样本", "最优测试准确率", "最佳轮次"],
        [
            ["全标签监督", total_train, 0, f"{acc_full:.2f}%", exps["full"]["metrics"].get("best_epoch", "-")],
            ["10%标签监督", labeled, "未使用", f"{acc_limited:.2f}%", exps["limited"]["metrics"].get("best_epoch", "-")],
        ],
        font_size=7.6,
    )
    add_image(doc, figs["acc"], "图2-1 精度对比与半监督恢复情况", width=6.25)
    add_para(doc, f"由表2-2和图2-1可知，标签从{total_train}张减少到{labeled}张后，测试准确率由{acc_full:.2f}%下降到{acc_limited:.2f}%，下降{drop:.2f}个百分点，说明标签不足会显著削弱模型的泛化能力。")

    add_heading(doc, "2.3 训练曲线、半监督方法与混淆矩阵", 1)
    add_image(doc, figs["curves"], "图2-2 三组实验训练/测试损失与准确率曲线", width=6.25)
    add_para(doc, "FixMatch通过伪标签和一致性正则化利用无标签数据。模型先对无标签图像的弱增强版本进行预测，当最大类别概率超过阈值τ=0.95时，将预测类别作为伪标签；随后对同一图像的强增强版本进行预测，并计算其与伪标签之间的一致性损失。总损失由有标签交叉熵损失和无标签一致性损失组成。")
    add_image(doc, figs["flow"], "图2-3 FixMatch半监督学习实现框图", width=6.25)
    add_image(doc, figs["cm"], "图2-4 三组实验归一化混淆矩阵", width=6.25)

    add_heading(doc, "2.4 结果分析与结论", 1)
    add_table(
        doc,
        "表2-3 加入半监督后的精度恢复情况",
        ["方法", "有标签样本", "无标签样本", "最优准确率", "相对10%监督变化"],
        [
            ["全标签监督", total_train, 0, f"{acc_full:.2f}%", f"+{drop:.2f}pp"],
            ["10%标签监督", labeled, 0, f"{acc_limited:.2f}%", "基准"],
            ["FixMatch半监督", labeled, unlabeled, f"{acc_fix:.2f}%", f"+{gain:.2f}pp"],
        ],
        font_size=7.6,
    )
    add_para(doc, f"在相同{labeled}张有标签样本条件下，FixMatch将测试准确率从{acc_limited:.2f}%提升到{acc_fix:.2f}%，提升{gain:.2f}个百分点；相对于标签减少造成的{drop:.2f}个百分点损失，恢复了约{recover:.2f}%的精度损失。该结果说明，无标签样本能够通过高置信度伪标签与一致性约束为模型提供额外训练信息。")
    add_para(doc, f"同时，FixMatch仍比全标签监督低{gap:.2f}个百分点，说明半监督学习能够缓解标签不足，但不能完全替代真实标签。总体而言，三组实验形成了“全标签最高、少标签明显下降、半监督部分恢复”的清晰结果，实验达到了实践二关于有限标签影响与半监督改进效果的验证目的。")

    add_heading(doc, "2.5 环境问题与解决过程", 1)
    add_para(doc, "实验环境为Python和PyTorch。数据集中图像尺寸不完全一致，因此在数据读取阶段统一转换为灰度图并缩放到128×128；为降低torchvision版本不匹配风险，代码中自定义ImageFolder和常用图像增强操作；少量标签训练样本过少时，半监督训练采用有标签样本重采样，并使用0.95置信度阈值筛选伪标签，以减少错误伪标签带来的干扰。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Sohn K, Berthelot D, Li C L, et al. FixMatch: Simplifying Semi-Supervised Learning with Consistency and Confidence. NeurIPS, 2020.",
        "[2] He K, Zhang X, Ren S, Sun J. Deep Residual Learning for Image Recognition. CVPR, 2016.",
    ]
    for ref in refs:
        add_para(doc, ref, indent=False, size=8.5)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)

# -----------------------------
# CLI
# -----------------------------

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Generate Practice 2 MSTAR semi-supervised learning report.")
    p.add_argument("--runs", default="runs", help="Directory containing experiment run folders")
    p.add_argument("--out", default="组号-人工智能实训II-实践2.docx", help="Output DOCX filename")
    p.add_argument("--fig-dir", default="practice2_figures", help="Output figure directory")
    p.add_argument("--group", default="组号", help="Group name; kept for filename/report metadata")
    p.add_argument("--no-docx", action="store_true", help="Only generate figures, not DOCX")
    return p.parse_args()


def main() -> None:
    args = parse_args()
    setup_matplotlib_font()

    runs_dir = Path(args.runs)
    fig_dir = Path(args.fig_dir)
    out_docx = Path(args.out)

    dirs = find_experiment_dirs(runs_dir)
    exps = {k: read_experiment(dirs[k]) for k in ORDER}

    print("Found experiments:")
    for k in ORDER:
        m = exps[k]["metrics"]
        print(f"  {k:8s}: {dirs[k]}  acc={pct(m.get('best_test_acc', 0)):.2f}%  epoch={m.get('best_epoch', '-')}")

    figs = {
        "acc": make_accuracy_recovery_figure(exps, fig_dir),
        "curves": make_training_curves_figure(exps, fig_dir),
        "cm": make_confusion_matrices_figure(exps, fig_dir),
        "flow": make_fixmatch_flowchart(fig_dir),
    }
    make_individual_curve_figures(exps, fig_dir)

    if not args.no_docx:
        build_report(exps, figs, out_docx, group_name=args.group)
        print(f"\nReport generated: {out_docx.resolve()}")
    print(f"Figures generated: {fig_dir.resolve()}")


if __name__ == "__main__":
    main()
