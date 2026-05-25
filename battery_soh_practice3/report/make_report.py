from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/mnt/data/battery_soh_practice3')
OUT = ROOT / 'outputs'
REPORT = ROOT / 'report' / '组长学号-人工智能实训II-实践3.docx'
MD = ROOT / 'report' / '组长学号-人工智能实训II-实践3.md'

metrics = pd.read_csv(OUT / 'metrics_summary.csv')
splits = pd.read_csv(OUT / 'split_summary.csv')
pcc = pd.read_csv(OUT / 'pcc_all_scenarios.csv')
features = pd.read_csv(OUT / 'features_all.csv')

scenario_cn = {
    'A_random_B0005_60_20_20': 'A：B0005 单电池随机 60/20/20 划分',
    'B_chrono_B0005_first60_last40': 'B：B0005 前 60% cycle 训练，后 40% cycle 测试',
    'C_transfer_B0005_to_B0007_target10': 'C：B0005 源电池 + B0007 前 10% 适配，B0007 后 90% 测试',
}


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = 'Noto Sans CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_from_df(doc, df, title=None, float_cols=None):
    if title:
        p = doc.add_paragraph(title)
        p.style = 'Caption'
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        set_cell_text(hdr[j], col, bold=True)
        set_cell_shading(hdr[j], 'D9EAF7')
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if float_cols and col in float_cols:
                try:
                    val = f'{float(val):.4f}'
                except Exception:
                    pass
            set_cell_text(cells[j], val)
    return table


def add_picture(doc, path, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.style = 'Caption'
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.15
    return p


def style_doc(doc):
    styles = doc.styles
    for style_name in ['Normal', 'Body Text']:
        style = styles[style_name]
        style.font.name = 'Noto Sans CJK SC'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
        style.font.size = Pt(10.5)
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
        style = styles[style_name]
        style.font.name = 'Noto Sans CJK SC'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    styles['Caption'].font.name = 'Noto Sans CJK SC'
    styles['Caption']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    styles['Caption'].font.size = Pt(9)


doc = Document()
style_doc(doc)
sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.85)
sec.right_margin = Inches(0.85)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('人工智能实训 II - 实践三\n')
r.bold = True
r.font.size = Pt(20)
r.font.name = 'Noto Sans CJK SC'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
r = p.add_run('基于深度学习的锂离子电池 SOH 预测实验报告')
r.bold = True
r.font.size = Pt(18)
r.font.name = 'Noto Sans CJK SC'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('文件名：组长学号-人工智能实训II-实践3\n').bold = True
p.add_run('组号：______    组长学号：______    姓名：______    日期：______')

doc.add_heading('摘要', level=1)
add_para(doc, '本实验围绕锂离子电池健康状态（State of Health, SOH）预测问题，使用 NASA Battery Aging ARC-FY08Q4 数据集构建了一个可复现实验流程。实验首先解析 .mat 原始文件，提取 discharge cycle 中的容量、端电压、电流、温度和时间序列；随后以额定容量 2Ah 为基准定义 SOH = Capacity / 2.0 × 100%，并从放电曲线中构造到达固定电压阈值时间、温度变化、曲线均值与斜率等健康特征。特征选择阶段使用训练集 Pearson Correlation Coefficient（PCC）评价特征与 SOH 的相关性，并选择高相关特征作为模型输入。模型采用 PyTorch 实现的多层感知机（MLP）回归网络，分别完成随机划分、时间顺序划分和跨电池少量目标数据适配三种实验设置。')
add_para(doc, '实验结果表明，基于放电曲线健康特征的 MLP 模型在三个场景下均取得较低预测误差，其中随机划分结果最好；按时间外推和跨电池测试更接近真实应用，难度更高，但仍保持较好的拟合效果。')

# 1

doc.add_heading('1. 实验目的与应用背景', level=1)
add_para(doc, '锂离子电池广泛用于便携电子设备、新能源汽车和储能系统。随着充放电循环增加，电池可用容量逐渐衰减，内阻、温度响应和电压平台也会发生变化。SOH 是电池管理系统（Battery Management System, BMS）中的关键状态变量，通常用于表征电池当前可用容量相对于额定容量或初始容量的退化程度。准确预测 SOH 有助于判断电池是否接近失效阈值，并支持维护、更换和安全管理。')
add_para(doc, '传统电化学模型或等效电路模型具有物理意义，但参数辨识较复杂；数据驱动方法能够利用历史传感器数据学习电压、电流、温度与容量衰减之间的非线性关系。考虑到本实验数据量较小、每个 discharge cycle 长度不一致，本实验不直接将原始曲线输入复杂序列模型，而是先构造具有物理含义的 cycle-level 健康特征，再使用轻量 MLP 模型进行回归预测。')

# 2

doc.add_heading('2. 数据集与任务定义', level=1)
add_para(doc, '本实验使用 NASA Prognostics Center of Excellence 发布的 Battery Aging ARC-FY08Q4 数据，包括 B0005、B0006、B0007 和 B0018 四个锂离子电池。每个文件包含 charge、discharge 和 impedance 三类操作记录。本实验聚焦 SOH 预测，因此仅使用 discharge cycle：该部分包含端电压、电流、温度、时间序列以及该次放电测得的 Capacity。')
add_para(doc, '根据数据说明，该组电池额定容量为 2Ah，实验终止条件为容量衰减 30%，即从 2Ah 下降到约 1.4Ah。因此本实验定义监督学习标签为：SOH(%) = Capacity / 2.0Ah × 100%。输入特征来自同一放电过程中的电压、电流、温度和时间曲线，不把 Capacity 本身作为输入。')
summary = features.groupby('battery_id').agg(discharge_cycles=('cycle_index','count'), min_soh=('soh_percent','min'), max_soh=('soh_percent','max')).reset_index()
summary.columns = ['电池编号', '放电 cycle 数', '最小 SOH(%)', '最大 SOH(%)']
add_table_from_df(doc, summary, '表1 预处理后各电池放电 cycle 数与 SOH 范围', float_cols=['最小 SOH(%)','最大 SOH(%)'])
add_picture(doc, OUT / 'capacity_degradation.png', '图1 四个电池的 SOH 衰减曲线')

# 3 preprocessing

doc.add_heading('3. 数据预处理与数据划分', level=1)
doc.add_heading('3.1 数据读取与异常剔除', level=2)
add_para(doc, '原始数据为 MATLAB .mat 结构体，顶层字段为 cycle。程序逐个读取所有 B*.mat 文件，只保留 type = discharge 的记录。对每条放电记录执行以下清洗：检查电压、电流、温度和时间序列长度一致；剔除非有限值；按照时间升序排列并去除重复时间点；剔除长度小于 10、持续时间非正、容量不在合理范围（0.5Ah, 2.2Ah] 内的样本。经处理后得到 636 条有效 discharge cycle 记录。')

doc.add_heading('3.2 特征工程', level=2)
add_para(doc, '为了使模型输入兼具可解释性与稳定性，本实验从每个 discharge cycle 中提取三类特征：第一类是放电曲线整体统计量，如放电持续时间、端电压均值/标准差、温度均值/最大值/范围、电压斜率等；第二类是健康指示特征，如电压降至 4.1V、4.0V、3.9V、3.8V、3.7V、3.6V、3.5V 的时间；第三类是固定相对时间点的曲线采样特征，如 25%、50%、75% 放电时间处的电压与温度。')
add_para(doc, '注意：Capacity 是监督标签，不能作为输入；此外，由电流对时间直接积分得到的 Ah 与 Capacity 高度等价，完整放电能量 energy_wh 也与容量强相关，容易使任务过于接近直接容量计算。因此代码中将 charge_ah_integral 和 energy_wh 均排除在默认模型输入之外。')

doc.add_heading('3.3 三种实验划分', level=2)
add_para(doc, '按照实践要求，本实验实现三种数据划分。A 场景在单个电池 B0005 上随机选择 60% 作为训练集、20% 作为验证集、20% 作为测试集，用于验证同分布情况下模型是否能学习放电特征与 SOH 的关系。B 场景在 B0005 上按 cycle 时间顺序划分，前 60% cycle 用于模型开发，后 40% cycle 用于测试，更接近“用早期寿命预测后期寿命”的实际问题。C 场景选择 B0005 作为源电池，B0007 前 10% cycle 加入训练/验证作为目标电池少量适配数据，B0007 后 90% cycle 用于测试，用于模拟跨电池泛化与迁移。')
splits_pivot = splits.copy()
splits_pivot['场景'] = splits_pivot['scenario'].map(scenario_cn)
splits_pivot = splits_pivot[['场景','split','n','batteries']]
splits_pivot.columns = ['场景','集合','样本数','电池']
add_table_from_df(doc, splits_pivot, '表2 三种实验场景的数据划分')

# 4 PCC

doc.add_heading('4. 特征选择：PCC 相关性分析', level=1)
add_para(doc, 'Pearson 相关系数用于衡量单个特征与连续标签 SOH 之间的线性相关程度，取值范围为 [-1, 1]。绝对值越大，说明该特征与 SOH 的单变量相关性越强。为避免测试集信息泄露，本实验在每个场景中只使用训练集计算 PCC，然后选择绝对值最高的前 8 个特征输入模型。')
for scen in metrics['scenario']:
    top = pcc[pcc['scenario']==scen].head(8)[['feature','pcc','abs_pcc']].copy()
    top.columns = ['特征', 'PCC', '|PCC|']
    add_table_from_df(doc, top, f'表 PCC-{list(metrics["scenario"]).index(scen)+1} {scenario_cn[scen]} 的前 8 个 PCC 特征', float_cols=['PCC','|PCC|'])
    add_picture(doc, OUT / f'pcc_{scen}.png', f'图 PCC-{list(metrics["scenario"]).index(scen)+1} {scenario_cn[scen]} 的 PCC 特征排序')

# 5 Model

doc.add_heading('5. 模型选择与训练方法', level=1)
add_para(doc, '本实验选择多层感知机（MLP）作为深度学习回归模型。主要理由如下：第一，经过特征工程后，每个样本已经是固定长度 tabular 特征，MLP 能直接处理；第二，NASA ARC-FY08Q4 每个电池的有效放电 cycle 数较少，直接训练 LSTM、Transformer 等序列模型容易过拟合；第三，MLP 结构简单、训练稳定，便于将性能变化归因到数据划分和特征质量。')
add_para(doc, '网络结构为 Linear(input_dim, 64) + ReLU + Dropout(0.05) + Linear(64, 32) + ReLU + Linear(32, 1)。输出为 SOH 百分比的标准化值，训练后再反标准化为 SOH(%)。优化器使用 AdamW，学习率为 0.002，损失函数为 SmoothL1Loss，并根据验证集损失进行早停。输入特征和输出标签均使用训练集统计量完成标准化，验证集和测试集只使用训练集拟合得到的转换器，避免数据泄露。')

# 6 Results

doc.add_heading('6. 模型训练与测试结果', level=1)
met = metrics.copy()
met['场景'] = met['scenario'].map(scenario_cn)
met = met[['场景','n_train','n_val','n_test','top_k_features','best_epoch','MAE','RMSE','R2']]
met.columns = ['场景','训练样本','验证样本','测试样本','PCC特征数','最佳epoch','MAE','RMSE','R2']
add_table_from_df(doc, met, '表3 三种场景的 SOH 测试结果（MAE/RMSE 单位为 SOH 百分点）', float_cols=['MAE','RMSE','R2'])
add_para(doc, '从结果看，随机划分场景 A 的 MAE 约为 0.10 个 SOH 百分点，说明在同一电池、同分布训练测试条件下，放电曲线特征与 SOH 的映射关系较容易学习。时间顺序划分场景 B 的误差上升到约 0.41 个 SOH 百分点，这是因为模型需要从早期循环外推到后期退化阶段，任务难度明显高于随机划分。跨电池场景 C 的误差约为 0.75 个 SOH 百分点，虽然加入的目标电池数据只有前 10%，但通过源电池训练和少量目标适配，模型仍能保持较高 R2。')
for scen in metrics['scenario']:
    add_picture(doc, OUT / f'predictions_{scen}.png', f'图 预测-{list(metrics["scenario"]).index(scen)+1} {scenario_cn[scen]} 测试集真实 SOH 与预测 SOH')

# 7 problems

doc.add_heading('7. 实验过程中出现的问题与解决', level=1)
add_para(doc, '问题一：原始 .mat 文件为多层嵌套结构，直接读取后字段访问复杂。解决方法是在 features.py 中封装 loadmat 与 discharge cycle 解析函数，将每个 cycle 统一转换为一行表格数据。')
add_para(doc, '问题二：不同 discharge cycle 的时间序列长度不一致，不能直接堆叠成矩阵。解决方法是提取固定长度的统计特征、阈值时间特征和固定比例时间点采样特征。')
add_para(doc, '问题三：Capacity 与由电流积分得到的 Ah 特征高度等价，完整放电能量 energy_wh 也会与 SOH 呈近乎单调关系。如果直接作为输入，模型性能会过于乐观。解决方法是在 model_feature_columns 中显式排除 capacity_ah、soh_percent、charge_ah_integral 和 energy_wh，仅保留可解释但不直接复制标签的曲线健康指标。')
add_para(doc, '问题四：时间顺序划分和跨电池测试比随机划分更难。解决方法是加入验证集早停，并在跨电池场景中使用目标电池前 10% cycle 作为少量适配数据，使模型获得目标电池初期退化水平的信息。')

# 8 conclusion

doc.add_heading('8. 结论', level=1)
add_para(doc, '本实验完成了从文献背景、数据预处理、特征工程、PCC 特征选择、PyTorch 模型训练到三种测试场景评估的完整流程。结果表明，基于放电曲线健康特征的 MLP 模型能够有效预测 NASA 锂离子电池 SOH；同时，随机划分结果容易偏乐观，按时间顺序外推和跨电池迁移更能反映实际应用难度。后续可进一步尝试增量容量（IC）曲线特征、基于早期充电片段的在线估计，以及更多源电池联合训练的迁移学习策略。')

# references

doc.add_heading('参考文献', level=1)
refs = [
    'NASA Prognostics Center of Excellence. Battery Data Set, NASA Ames Research Center. https://www.nasa.gov/intelligent-systems-division/discovery-and-systems-health/pcoe/pcoe-data-set-repository/',
    'NASA Open Data Portal. Li-ion Battery Aging Datasets. https://data.nasa.gov/dataset/li-ion-battery-aging-datasets',
    'B. Saha and K. Goebel. Battery Data Set, NASA Prognostics Data Repository, NASA Ames Research Center, Moffett Field, CA, 2007.',
    'D. Roman, S. Saxena, V. Robu, M. Pecht, and D. Flynn. Machine learning pipeline for battery state of health estimation, 2021.',
    'Z. Nie, J. Zhao, Q. Li, and Y. Qin. CyFormer: Accurate State-of-Health Prediction of Lithium-Ion Batteries via Cyclic Attention, 2023.',
    'Y. Lian et al. A Novel Capacity Estimation Method for Lithium-Ion Batteries Based on the Adam Algorithm, Batteries, 2025.'
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f'[{i}] {ref}')
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)

# Appendix

doc.add_page_break()
doc.add_heading('附录：源代码文件说明', level=1)
code_table = pd.DataFrame([
    ['run_all.py', '一键运行完整实验：数据提取、特征表生成、三种划分训练、指标与图表输出。'],
    ['src/features.py', '读取 NASA .mat 数据，提取 discharge cycle 特征，定义模型输入列。'],
    ['src/splits.py', '实现 A/B/C 三种数据划分方案。'],
    ['src/model.py', '定义 PyTorch MLP SOH 回归模型。'],
    ['src/train_eval.py', '实现 PCC 特征选择、标准化、训练、早停与 MAE/RMSE/R2 评估。'],
    ['outputs/*.csv', '实验记录数据，包括特征表、划分表、PCC 结果、测试集预测结果和指标汇总。'],
], columns=['文件', '说明'])
add_table_from_df(doc, code_table, '表4 项目文件说明')

doc.save(REPORT)

# Markdown copy
md_lines = []
md_lines.append('# 人工智能实训 II - 实践三：基于深度学习的锂离子电池 SOH 预测\n')
md_lines.append('本报告的完整排版版本见同目录 DOCX 文件。\n')
md_lines.append('## 测试结果\n')
md_lines.append(metrics.to_markdown(index=False))
md_lines.append('\n## 文件说明\n源代码在项目根目录和 src/ 目录中，实验记录在 outputs/ 目录中。\n')
MD.write_text('\n'.join(md_lines), encoding='utf-8')
print(REPORT)
