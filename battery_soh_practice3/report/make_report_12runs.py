from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/mnt/data/battery_soh_practice3')
OUT = ROOT / 'outputs'
REPORT = ROOT / 'report' / '实践三-基于深度学习的锂离子电池SOH预测实验报告-12次独立实验.docx'

metrics = pd.read_csv(OUT / 'metrics_summary.csv')
splits = pd.read_csv(OUT / 'split_summary.csv')
pcc = pd.read_csv(OUT / 'pcc_all_scenarios.csv')
features = pd.read_csv(OUT / 'features_all.csv')

case_name = {
    'A': 'A：单电池随机 60%/20%/20% 划分',
    'B': 'B：单电池前 60% cycle 训练，后 40% cycle 测试',
    'C': 'C：单源电池训练 + 目标电池前 10% cycle 适配，后 90% 测试',
}

def scenario_label(row):
    if row['case'] == 'A':
        return f"A-{row['target_battery']}"
    if row['case'] == 'B':
        return f"B-{row['target_battery']}"
    return f"C-{row['source_battery']}→{row['target_battery']}"

metrics['场景编号'] = metrics.apply(scenario_label, axis=1)

# ---------- helpers ----------
def set_run_font(run, size=10.5, bold=False):
    run.font.name = 'Noto Sans CJK SC'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text, bold=False, size=8.5, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, 14 if level == 1 else 12, True)
    return p


def add_table(doc, df, caption, float_cols=None, font_size=8.0):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, 9, True)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        set_cell_text(hdr[j], col, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr[j], 'D9EAF7')
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if float_cols and col in float_cols:
                try:
                    val = f'{float(val):.4f}'
                except Exception:
                    pass
            set_cell_text(cells[j], val, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    return table


def add_picture(doc, path, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, 9, True)


def style_doc(doc):
    styles = doc.styles
    for style_name in ['Normal', 'Body Text']:
        st = styles[style_name]
        st.font.name = 'Noto Sans CJK SC'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
        st.font.size = Pt(10.5)
    for name in ['Heading 1', 'Heading 2', 'Heading 3']:
        st = styles[name]
        st.font.name = 'Noto Sans CJK SC'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')


def set_margins(sec):
    sec.top_margin = Cm(1.9)
    sec.bottom_margin = Cm(1.9)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

# ---------- document ----------
doc = Document()
style_doc(doc)
set_margins(doc.sections[0])

# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('实践三\n')
set_run_font(r, 20, True)
r = p.add_run('基于深度学习的锂离子电池 SOH 预测实验报告\n')
set_run_font(r, 18, True)
r = p.add_run('（A/B/C 三种划分均对 B0005、B0006、B0007、B0018 独立训练与测试）')
set_run_font(r, 12, True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for text in ['课程：人工智能开发实训 II\n', '组长学号：__________    姓名：__________    组号：__________\n', '日期：__________']:
    r = p.add_run(text)
    set_run_font(r, 11, False)

doc.add_paragraph()

add_heading(doc, '1. 实验目的与任务理解', 1)
add_para(doc, '本实践任务要求使用 PyTorch 深度学习方法解决锂离子电池健康状态（State of Health, SOH）预测问题。实验需要完成应用背景阅读、数据读取与预处理、特征选择、模型选择、模型训练与测试，并使用 MAE、RMSE、R2 等指标评价预测效果。')
add_para(doc, '根据重新梳理后的任务要求，本实验不再只选取一个电池或一个跨电池组合进行演示，而是对 NASA 数据集中的 B0005、B0006、B0007、B0018 四个电池分别执行 A、B、C 三种实验设置。也就是说，A 场景训练 4 次、测试 4 次；B 场景训练 4 次、测试 4 次；C 场景训练 4 次、测试 4 次，共完成 12 次独立训练与测试。')

add_heading(doc, '2. 应用背景与 SOH 定义', 1)
add_para(doc, '锂离子电池在新能源汽车、储能系统和便携式电子设备中广泛使用。随着充放电循环增加，电池容量逐渐衰减，电压平台、温度响应和内阻状态也会发生变化。SOH 是电池管理系统中的重要状态量，通常反映当前容量相对于额定容量或初始容量的保持程度。准确预测 SOH 有助于判断电池寿命阶段、制定维护策略并降低安全风险。')
add_para(doc, '本实验使用 NASA Battery Aging ARC-FY08Q4 数据集，其中 B0005、B0006、B0007、B0018 四个 .mat 文件均包含 charge、discharge、impedance 三类 cycle。SOH 预测以 discharge cycle 为核心，因为 discharge 记录中包含实际测得的容量 Capacity。数据说明中电池额定容量为 2Ah，失效条件约为容量衰减到 1.4Ah，因此本实验将标签定义为 SOH(%) = Capacity / 2.0Ah × 100%。')

summary = features.groupby('battery_id').agg(
    discharge_cycles=('cycle_index', 'count'),
    min_soh=('soh_percent', 'min'),
    max_soh=('soh_percent', 'max'),
    min_capacity=('capacity_ah', 'min'),
    max_capacity=('capacity_ah', 'max')
).reset_index()
summary.columns = ['电池编号', '有效放电cycle数', '最小SOH(%)', '最大SOH(%)', '最小容量(Ah)', '最大容量(Ah)']
add_table(doc, summary, '表 3-1 预处理后四个电池的样本数量与容量范围', float_cols=['最小SOH(%)', '最大SOH(%)', '最小容量(Ah)', '最大容量(Ah)'])
add_picture(doc, OUT / 'capacity_degradation.png', '图 3-1 四个电池的 SOH 衰减曲线', width=6.3)

add_heading(doc, '3. 数据预处理与特征工程', 1)
add_heading(doc, '3.1 数据读取与异常剔除', 2)
add_para(doc, '原始数据为 MATLAB .mat 结构体，程序通过 scipy.io.loadmat 读取每个电池文件，并逐个解析 cycle。由于本任务目标是 SOH 预测，实验仅保留 type = discharge 的记录。对于每条 discharge cycle，读取 Time、Voltage_measured、Current_measured、Temperature_measured 和 Capacity 字段。')
add_para(doc, '预处理阶段主要包括：剔除长度过短的曲线；剔除时间、电压、电流、温度中存在非有限值的记录；按时间升序排序并去除重复时间点；剔除持续时间非正或 Capacity 不在合理范围内的记录。最终四个电池共得到 636 条有效 discharge cycle 样本。')

add_heading(doc, '3.2 输入特征构造', 2)
add_para(doc, '不同 discharge cycle 的原始时间序列长度并不完全相同，不能直接拼接成固定长度矩阵。因此，本实验将每个 cycle 转换为固定长度的表格特征。构造的特征主要包括三类：一是整体统计特征，例如放电持续时间、端电压均值、端电压标准差、温度均值、温度范围、电压斜率等；二是阈值时间特征，例如电压首次下降到 4.1V、4.0V、3.9V、3.8V、3.7V、3.6V、3.5V 所需时间；三是固定比例时间点特征，例如 25%、50%、75% 放电时间处的电压和温度。')
add_para(doc, '为避免标签泄露，Capacity 只作为监督标签，不作为模型输入。此外，由电流对时间积分得到的 Ah 特征与 Capacity 近似等价，完整放电能量 energy_wh 也与容量强相关，容易使任务变成间接容量计算。因此默认模型输入中排除了 capacity_ah、soh_percent、charge_ah_integral 和 energy_wh。')

add_heading(doc, '4. 三种数据划分方案', 1)
add_para(doc, '本次修改后的关键点是：A、B、C 三种划分都必须对四个电池单独执行，不能把四个电池混在一起形成一个总实验。具体设置如下。')
add_para(doc, 'A 场景：对单个电池的全部有效 discharge cycle 随机划分，60% 作为训练集，20% 作为验证集，20% 作为测试集。B0005、B0006、B0007、B0018 分别独立运行一次。')
add_para(doc, 'B 场景：对单个电池按 cycle 时间顺序划分，前 60% cycle 用于模型开发，后 40% cycle 用于测试。前 60% 中再划出一小部分作为验证集用于早停。四个电池分别独立运行一次。')
add_para(doc, 'C 场景：每次只选择一个源电池，不合并多个源电池；目标电池前 10% cycle 用作少量目标域适配，目标电池后 90% cycle 用作测试。为了让四个目标电池都被单独测试一次，本实验设置四个 C 场景配对：B0007→B0005、B0007→B0006、B0005→B0007、B0005→B0018。')

split_table = metrics[['case','source_battery','target_battery','n_train','n_val','n_test']].copy()
split_table['场景'] = metrics['场景编号']
split_table = split_table[['场景','case','source_battery','target_battery','n_train','n_val','n_test']]
split_table.columns = ['场景', '划分类型', '源电池', '目标/测试电池', '训练样本', '验证样本', '测试样本']
add_table(doc, split_table, '表 3-2 A/B/C 三种设置下的 12 次独立数据划分', font_size=7.8)

add_heading(doc, '5. PCC 特征选择', 1)
add_para(doc, 'Pearson Correlation Coefficient（PCC）用于衡量单个特征与连续标签 SOH 之间的线性相关程度。PCC 的取值范围为 [-1, 1]，绝对值越大表示该特征与 SOH 的单变量相关性越强。为避免测试集信息泄露，每个场景都只在对应训练集上计算 PCC，并选择 |PCC| 排名前 8 的特征作为 MLP 输入。')
# top 3 pcc features per scenario for readability
pcc_rows = []
for _, row in metrics.iterrows():
    scen = row['scenario']
    sub = pcc[pcc['scenario'] == scen].sort_values('abs_pcc', ascending=False).head(3)
    feat_desc = '; '.join([f"{r.feature}({r.pcc:.3f})" for r in sub.itertuples()])
    pcc_rows.append({'场景': row['场景编号'], '前3个PCC特征（括号内为PCC）': feat_desc})
pcc_df = pd.DataFrame(pcc_rows)
add_table(doc, pcc_df, '表 3-3 各独立实验训练集上的代表性 PCC 特征', font_size=7.2)
add_para(doc, '从 PCC 结果看，time_to_3p5v_s、time_to_3p6v_s、time_to_3p7v_s、duration_s、early_voltage_slope 等特征频繁出现在前列。这符合电池退化机理：随着容量下降，在相同放电条件下电压平台持续时间缩短，达到低电压阈值的时间提前；同时放电早期电压下降斜率、放电持续时间等曲线形态指标也会随健康状态变化。因此，基于 PCC 选择上述曲线健康特征具有合理性。')

add_heading(doc, '6. 深度学习模型选择与训练方法', 1)
add_para(doc, '本实验采用 PyTorch 实现多层感知机（MLP）回归模型。选择 MLP 的原因是：经过特征工程后，每个 discharge cycle 已经被表示为固定长度的表格特征；NASA 数据集中每个电池有效 cycle 数较少，直接训练 LSTM 或 Transformer 等较复杂序列模型容易过拟合；MLP 结构简单、训练稳定，便于对比不同划分策略对 SOH 预测效果的影响。')
add_para(doc, '模型结构为 Linear(input_dim, 64) + ReLU + Dropout(0.05) + Linear(64, 32) + ReLU + Linear(32, 1)。训练时使用 AdamW 优化器，学习率为 0.002，损失函数为 SmoothL1Loss。输入特征使用训练集的中位数补全缺失值，并用训练集统计量进行标准化；SOH 标签也进行标准化，预测后再反标准化为 SOH 百分比。模型根据验证集损失早停，最多训练 600 个 epoch。')

add_heading(doc, '7. 模型测试结果与分析', 1)
res_table = metrics[['场景编号','source_battery','target_battery','n_train','n_val','n_test','MAE','RMSE','R2']].copy()
res_table.columns = ['场景','源电池','目标/测试电池','训练样本','验证样本','测试样本','MAE','RMSE','R2']
add_table(doc, res_table, '表 3-4 12 次独立训练与测试的 SOH 预测结果', float_cols=['MAE','RMSE','R2'], font_size=7.5)

summary_case = metrics.groupby('case').agg(MAE均值=('MAE','mean'), RMSE均值=('RMSE','mean'), R2均值=('R2','mean'), MAE最小=('MAE','min'), MAE最大=('MAE','max')).reset_index()
summary_case['实验设置'] = summary_case['case'].map(case_name)
summary_case = summary_case[['实验设置','MAE均值','RMSE均值','R2均值','MAE最小','MAE最大']]
add_table(doc, summary_case, '表 3-5 按 A/B/C 汇总的平均测试性能', float_cols=['MAE均值','RMSE均值','R2均值','MAE最小','MAE最大'], font_size=8.0)

add_picture(doc, OUT / 'metrics_12_runs.png', '图 3-2 12 次独立实验的 MAE、RMSE 与 R2 对比', width=6.3)
add_picture(doc, OUT / 'predictions_A_four_batteries.png', '图 3-3 A 场景四个电池的测试集真实 SOH 与预测 SOH', width=6.3)
add_picture(doc, OUT / 'predictions_B_four_batteries.png', '图 3-4 B 场景四个电池的测试集真实 SOH 与预测 SOH', width=6.3)
add_picture(doc, OUT / 'predictions_C_four_batteries.png', '图 3-5 C 场景四个目标电池的测试集真实 SOH 与预测 SOH', width=6.3)

add_para(doc, f"从表 3-4 和表 3-5 可以看到，A 场景平均 MAE 为 {metrics[metrics.case=='A'].MAE.mean():.3f} 个 SOH 百分点，平均 R2 为 {metrics[metrics.case=='A'].R2.mean():.4f}，说明在同一电池随机划分的同分布条件下，模型能够较准确地学习曲线特征与 SOH 的对应关系。")
add_para(doc, f"B 场景平均 MAE 上升到 {metrics[metrics.case=='B'].MAE.mean():.3f}，平均 R2 为 {metrics[metrics.case=='B'].R2.mean():.4f}。该场景要求模型用早期 cycle 外推后期退化阶段，比随机划分更接近实际预测任务，因此误差更高。其中 B0006 的时间外推效果最差，说明不同电池的退化轨迹和后期衰减形态存在差异。")
add_para(doc, f"C 场景平均 MAE 为 {metrics[metrics.case=='C'].MAE.mean():.3f}，平均 R2 为 {metrics[metrics.case=='C'].R2.mean():.4f}。该结果说明，在只使用一个源电池且目标电池只有前 10% cycle 可用的情况下，模型仍能获得一定跨电池泛化能力。但 B0006 作为目标电池时误差相对较高，提示跨电池 SOH 预测受到电池个体差异、初始容量差异和衰减速度差异影响。")

add_heading(doc, '8. 实验过程中出现的问题与解决过程', 1)
add_para(doc, '问题一：原始 .mat 文件为多层嵌套结构，直接读取后字段访问较复杂。解决方法是在 features.py 中封装读取函数，将每个 discharge cycle 统一解析为一行表格数据，便于后续特征工程和训练。')
add_para(doc, '问题二：不同 cycle 的时间序列长度不一致，无法直接作为普通神经网络输入。解决方法是提取固定长度的 cycle-level 特征，包括阈值时间、统计特征和固定比例时间点采样特征。')
add_para(doc, '问题三：Capacity 是标签，charge_ah_integral 和 energy_wh 又与容量高度相关。如果直接输入模型，测试结果会过于乐观。解决方法是在特征列定义中显式排除这些可能造成标签泄露或近似标签泄露的变量。')
add_para(doc, '问题四：最初实验只对 B0005 做 A/B，并只做了一个 C 场景，不能满足“四个电池都要单独跑 A/B/C”的要求。解决方法是重写 run_all.py 中的场景生成逻辑，分别为 B0005、B0006、B0007、B0018 构造 A、B、C 三类实验，共 12 个 scenario，每个 scenario 都独立划分、独立训练、独立测试。')
add_para(doc, '问题五：B 场景和 C 场景比 A 场景更难，部分目标电池误差明显增大。解决方法是保留验证集早停、使用训练集标准化参数、防止数据泄露，并在结果分析中单独解释时间外推和跨电池泛化的难度，而不是只报告随机划分下较高的指标。')

add_heading(doc, '9. 结论', 1)
add_para(doc, '本实验完成了基于 NASA 锂离子电池老化数据集的 SOH 预测实践。实验通过 discharge cycle 数据读取、异常剔除、健康特征构造、PCC 特征选择和 PyTorch MLP 回归模型，完成了 A、B、C 三类数据划分下的 12 次独立训练与测试。')
add_para(doc, '结果表明，随机划分 A 场景效果最好，但容易反映同分布拟合能力；时间顺序划分 B 场景和跨电池少量适配 C 场景更接近真实应用，预测难度明显更高。整体来看，放电曲线中的阈值时间、放电持续时间和早期电压斜率等特征能够有效表征容量衰减过程，结合轻量 MLP 模型可以实现较稳定的 SOH 回归预测。')

add_heading(doc, '参考文献', 1)
refs = [
    'NASA Prognostics Center of Excellence. Battery Data Set, NASA Ames Research Center.',
    'NASA Open Data Portal. Li-ion Battery Aging Datasets.',
    'Saha B., Goebel K. Battery Data Set, NASA Prognostics Data Repository, NASA Ames Research Center, Moffett Field, CA, 2007.',
    'Roman D., Saxena S., Robu V., Pecht M., Flynn D. Machine learning pipeline for battery state of health estimation, 2021.',
    'Nie Z., Zhao J., Li Q., Qin Y. CyFormer: Accurate State-of-Health Prediction of Lithium-Ion Batteries via Cyclic Attention, 2023.'
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    r = p.add_run(f'[{i}] {ref}')
    set_run_font(r, 9)

# Appendix
add_heading(doc, '附录：源代码与记录数据说明', 1)
code_table = pd.DataFrame([
    ['run_all.py', '一键运行完整实验。新版代码会为 A/B/C 三种设置分别生成四个电池的独立 scenario，共 12 次训练与测试。'],
    ['src/features.py', '读取 .mat 文件，提取 discharge cycle 特征，并定义排除标签泄露后的模型输入列。'],
    ['src/splits.py', '实现 A 随机划分、B 时间顺序划分、C 单源电池到目标电池前 10% 适配的划分函数。'],
    ['src/model.py', '定义 PyTorch MLP SOH 回归模型。'],
    ['src/train_eval.py', '实现 PCC 特征选择、缺失值补全、标准化、训练、早停和 MAE/RMSE/R2 评估。'],
    ['outputs/features_all.csv', '预处理后的全部 cycle-level 特征表。'],
    ['outputs/metrics_summary.csv', '12 次独立实验的测试指标汇总。'],
    ['outputs/split_summary.csv', '12 次独立实验的训练/验证/测试划分记录。'],
    ['outputs/pcc_all_scenarios.csv', '每个场景训练集上的 PCC 特征排序结果。'],
    ['outputs/predictions_all_scenarios.csv', '每个场景测试集上的真实 SOH、预测 SOH 和绝对误差。'],
], columns=['文件', '说明'])
add_table(doc, code_table, '表 3-6 源代码与记录数据文件说明', font_size=8.0)

doc.save(REPORT)
print(REPORT)
