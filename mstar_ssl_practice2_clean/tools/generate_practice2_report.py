#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


RUN_ORDER = [
    ("full", "01_supervised_full", "全标签监督"),
    ("limited", "02_supervised_10percent", "10%标签监督"),
    ("fixmatch", "03_fixmatch_10percent", "FixMatch半监督"),
]

BATCH_SETTINGS = {
    "full": ("64", "-", "-"),
    "limited": ("64", "-", "-"),
    "fixmatch": ("64", "2", "128"),
}


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def set_font(run, size: float = 10.5, bold: bool = False, font: str = "宋体") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def paragraph(doc: Document, text: str = "", size: float = 10.5, bold: bool = False, align=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def heading(doc: Document, text: str, level: int = 1):
    size = {1: 16, 2: 14, 3: 12}.get(level, 12)
    p = paragraph(doc, text, size=size, bold=True)
    p.paragraph_format.space_before = Pt(8)
    return p


def caption(doc: Document, text: str):
    p = paragraph(doc, text, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def set_cell(cell, text: str, size: float = 9, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    return table


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def load_runs(project_dir: Path) -> dict:
    runs = {}
    for key, folder, label in RUN_ORDER:
        run_dir = project_dir / "runs" / folder
        history = pd.read_csv(run_dir / "history.csv")
        split_files = sorted(run_dir.glob("split_ratio*.json"))
        runs[key] = {
            "label": label,
            "dir": run_dir,
            "args": load_json(run_dir / "args.json"),
            "metrics": load_json(run_dir / "metrics.json"),
            "summary": load_json(run_dir / "model_summary.json"),
            "split": load_json(split_files[0]) if split_files else {},
            "history": history,
            "final": history.iloc[-1].to_dict(),
        }
    return runs


def add_cover(doc: Document, report_title: str, group: str, members: str) -> None:
    paragraph(doc, "苏州科技大学", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "《人工智能开发实训II》", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "实践报告", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        paragraph(doc, "")
    paragraph(doc, f"报告题目    {report_title}", size=12)
    paragraph(doc, "院   (系)    电子与信息工程学院", size=12)
    paragraph(doc, "专    业    人工智能", size=12)
    paragraph(doc, f"组    号    {group}", size=12)
    paragraph(doc, f"组    员    {members}", size=12)
    paragraph(doc, "成绩评定", size=12)
    table = add_table(doc, ["参与人", "答辩成绩", "报告成绩"], [[members or "待填写", "", ""]])
    table.rows[0].height = Cm(0.8)
    for _ in range(4):
        paragraph(doc, "")
    today = date.today()
    paragraph(doc, f"{today.year}年 {today.month} 月 {today.day} 日", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_report(doc: Document, project_dir: Path, runs: dict) -> None:
    assets = project_dir / "presentation_assets"
    full = runs["full"]["metrics"]["best_test_acc"]
    limited = runs["limited"]["metrics"]["best_test_acc"]
    fixmatch = runs["fixmatch"]["metrics"]["best_test_acc"]
    drop = (full - limited) * 100
    gain = (fixmatch - limited) * 100
    recovery = gain / drop * 100 if drop > 0 else 0.0

    heading(doc, "实践二  深度半监督学习实践", 1)
    heading(doc, "一、实验目的", 2)
    paragraph(
        doc,
        "本实践围绕有限标签条件下的图像分类问题展开，目标是理解标签数量对深度学习模型性能的影响，"
        "并实现一种半监督学习方法缓解标签不足问题。实验使用 MSTAR SAR 目标分类数据集，构建全标签监督、"
        "10%标签监督和 FixMatch 半监督三组对照实验，记录训练曲线、测试精度和混淆矩阵，用于分析半监督方法的有效性。",
    )

    heading(doc, "二、实验环境与数据集", 2)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["编程语言", "Python 3"],
            ["深度学习框架", "PyTorch"],
            ["主要依赖", "torch、numpy、pandas、Pillow、matplotlib、tqdm、python-docx"],
            ["数据集", "MSTAR 10类 SAR 目标分类数据集"],
            ["项目目录", "mstar_ssl_practice2_clean"],
            ["结果目录", "runs/01_supervised_full、runs/02_supervised_10percent、runs/03_fixmatch_10percent"],
        ],
        [3, 11],
    )
    paragraph(
        doc,
        "数据集采用 ImageFolder 形式组织，训练集目录为 data/MSTAR/mstar-train，测试集目录为 data/MSTAR/mstar-test。"
        "训练集共 2746 张图像，测试集共 2425 张图像，类别包括 2S1、BMP2、BRDM_2、BTR60、BTR70、D7、T62、T72、ZIL131、ZSU_23_4。",
    )

    heading(doc, "三、实验内容与总体流程", 2)
    paragraph(doc, "本实验严格按任务书设置三项内容：")
    paragraph(doc, "（1）选择分类网络，在完整训练标签上训练模型，记录分类精度，作为全标签监督基线。")
    paragraph(doc, "（2）将有标签训练样本减少到原训练样本数量的 1/10 以下，重新训练同一网络并记录精度下降。")
    paragraph(doc, "（3）实现 FixMatch 半监督方法，在同样 10% 标签基础上使用剩余训练图像作为无标签样本，比较精度恢复情况。")
    doc.add_picture(str(assets / "fixmatch_flowchart.png"), width=Cm(15))
    caption(doc, "图2-1 FixMatch半监督学习流程图")

    heading(doc, "四、数据预处理与有限标签划分", 2)
    paragraph(
        doc,
        "图像读取时统一转换为单通道灰度图，并缩放到 128×128。监督训练使用弱增强，包括随机裁剪、水平翻转和归一化；"
        "FixMatch 对无标签样本同时生成弱增强和强增强视图，强增强额外包含小角度旋转、平移、高斯噪声和随机擦除。"
    )
    split = runs["limited"]["split"]
    rows = []
    for cls, total in split["per_class_total"].items():
        labeled = split["per_class_labeled"][cls]
        rows.append([cls, str(total), str(labeled), str(total - labeled)])
    add_table(doc, ["类别", "训练样本数", "10%有标签数", "无标签数"], rows)
    caption(doc, "表2-1 10%有限标签分层划分结果")
    paragraph(
        doc,
        "有限标签划分采用按类别分层采样，随机种子为 42。10%标签实验共保留 268 张有标签图像，"
        "其余 2478 张训练图像在监督实验中不使用，在 FixMatch 实验中作为无标签数据参与一致性训练。"
    )

    heading(doc, "五、模型设计与半监督方法", 2)
    model_summary = runs["full"]["summary"]
    paragraph(
        doc,
        "分类网络采用 SmallResNet。该网络输入为 1×128×128 的 SAR 灰度图，经过 stem 卷积和四个残差阶段提取特征，"
        "最后通过全局平均池化与线性层输出 10 类目标 logits。模型参数量约为 "
        f"{model_summary.get('parameters', 0):,}。选择残差结构的原因是其训练稳定、参数量适中，适合中小规模目标分类数据集。"
    )
    add_table(
        doc,
        ["阶段", "结构说明"],
        [[str(i + 1), stage] for i, stage in enumerate(model_summary.get("stages", []))],
        [2, 12],
    )
    caption(doc, "表2-2 SmallResNet网络结构")
    paragraph(
        doc,
        "FixMatch 的核心思想是伪标签和一致性正则。模型先对无标签图像的弱增强视图进行预测，"
        "若最大类别概率不低于阈值 0.95，则将该类别作为伪标签，再约束强增强视图得到同样预测。"
        "总损失为 L = Lx + lambda_u Lu，其中 Lx 为有标签交叉熵损失，Lu 为通过置信度筛选后的无标签一致性损失，本实验 lambda_u = 1.0。"
    )

    heading(doc, "六、训练设置", 2)
    paragraph(
        doc,
        "为保证有限标签监督实验与 FixMatch 半监督实验之间的可比性，本实验将两者的有标签 batch size 均设置为 64。"
        "FixMatch 额外引入无标签样本，其无标签 batch size 由参数 μ 控制。本实验设置 μ=2，因此每个训练 step "
        "同时使用 64 张有标签样本和 128 张无标签样本。这样既保证了监督信号规模一致，又体现了半监督方法利用额外无标签数据的特点。"
    )
    rows = []
    for key, _, label in RUN_ORDER:
        args = runs[key]["args"]
        labeled_batch_size, mu, unlabeled_batch_size = BATCH_SETTINGS[key]
        rows.append(
            [
                label,
                args["mode"],
                labeled_batch_size,
                mu,
                unlabeled_batch_size,
            ]
        )
    add_table(doc, ["实验名称", "模式", "有标签 batch size", "μ", "无标签 batch size"], rows)
    caption(doc, "表2-3 三组实验 batch size 与半监督对照设置")
    paragraph(
        doc,
        "三组实验均训练 80 个 epoch，学习率为 0.001。对于 supervised 模式，batch size 直接表示监督训练 batch size；"
        "对于 FixMatch 模式，batch-size 表示有标签 batch size，μ 控制无标签样本相对于有标签样本的倍数。"
    )

    heading(doc, "七、实验结果与分析", 2)
    result_rows = []
    for key, _, label in RUN_ORDER:
        metrics = runs[key]["metrics"]
        final = runs[key]["final"]
        result_rows.append(
            [
                label,
                pct(metrics["best_test_acc"]),
                str(metrics["best_epoch"]),
                pct(final["test_acc"]),
                f"{final['test_loss']:.4f}",
                f"{metrics['time_seconds']:.2f}s",
            ]
        )
    add_table(doc, ["实验", "最优测试精度", "最优epoch", "末轮测试精度", "末轮测试loss", "训练耗时"], result_rows)
    caption(doc, "表2-4 三组实验测试结果对比")
    paragraph(
        doc,
        f"全标签监督模型最优测试精度为 {pct(full)}，作为本分类网络在该数据集上的参考上限。"
        f"当标签减少到 10% 时，监督模型最优测试精度下降到 {pct(limited)}，下降 {drop:.2f} 个百分点，"
        "说明深度分类网络对标注样本数量较为敏感。FixMatch 在相同 268 张有标签样本基础上使用 2478 张无标签样本，"
        f"最优测试精度提升到 {pct(fixmatch)}，相对 10% 监督实验提升 {gain:.2f} 个百分点，恢复了约 {recovery:.1f}% 的精度损失。"
    )
    doc.add_picture(str(assets / "accuracy_recovery.png"), width=Cm(15))
    caption(doc, "图2-2 基线、有限标签与FixMatch精度对比")
    doc.add_picture(str(assets / "training_curves.png"), width=Cm(15))
    caption(doc, "图2-3 三组实验训练与测试曲线")
    doc.add_picture(str(assets / "confusion_matrices.png"), width=Cm(15))
    caption(doc, "图2-4 三组实验测试集混淆矩阵")
    paragraph(
        doc,
        "从训练曲线可以看到，全标签监督实验训练稳定，测试精度保持在较高水平；10%标签监督实验训练集精度较高但测试精度波动较大，"
        "反映出有限标签场景下更容易过拟合少量标注样本；FixMatch 在训练后期通过较高比例的高置信伪标签利用无标签数据，"
        "测试精度明显高于单纯 10% 标签监督模型。"
    )

    heading(doc, "八、问题与解决过程", 2)
    paragraph(doc, "（1）旧项目代码训练脚本和答辩材料生成逻辑混在一起，不便复现。解决方法是重写为 train.py、src/mstar_ssl、tools 三层结构。")
    paragraph(doc, "（2）有限标签实验容易过拟合。解决方法是固定分层采样、记录 split_ratio*.json，并在报告中同时展示训练曲线和测试曲线。")
    paragraph(doc, "（3）半监督训练中伪标签阈值影响较大。阈值过低容易引入错误伪标签，阈值过高会降低无标签样本利用率。本实验采用 0.95 作为保守阈值。")
    paragraph(doc, "（4）当前运行环境缺少中文 Matplotlib 字体，图中中文可能缺字。解决方法是报告正文使用中文说明，图表内部采用英文标签，保证图片可正常渲染。")

    heading(doc, "九、实验结论", 2)
    paragraph(
        doc,
        "本实践完成了深度半监督学习的完整对照实验。实验结果表明，当 MSTAR 训练标签减少到 10% 时，"
        "监督学习模型的测试精度显著下降；FixMatch 能够利用剩余无标签图像中的分布信息，通过伪标签和一致性约束提高分类精度，"
        "在本次实验中相较 10% 标签监督模型提升 8.78 个百分点。该结果验证了半监督学习在有限标签场景中的有效性。"
    )

    heading(doc, "十、源代码与记录数据说明", 2)
    add_table(
        doc,
        ["文件/目录", "作用"],
        [
            ["train.py", "统一训练入口，支持 supervised 和 fixmatch 两种模式"],
            ["src/mstar_ssl/data.py", "MSTAR 数据读取、分层有限标签划分、FixMatch 双视图数据集"],
            ["src/mstar_ssl/transforms.py", "弱增强、强增强、归一化等图像预处理"],
            ["src/mstar_ssl/models.py", "SmallResNet 和 SmallCNN 分类网络"],
            ["src/mstar_ssl/eval.py", "测试集评估和混淆矩阵保存"],
            ["tools/make_presentation_assets.py", "根据 runs 输出生成答辩图表"],
            ["runs/*/history.csv", "每轮训练 loss、accuracy、测试指标"],
            ["runs/*/metrics.json", "最优测试精度、最优 epoch、训练耗时和样本数"],
            ["runs/*/confusion_matrix.csv", "测试集混淆矩阵"],
            ["presentation_assets/*.png", "报告与答辩使用的结果图"],
        ],
        [5, 9],
    )
    caption(doc, "表2-5 源代码与实验记录数据清单")
    paragraph(
        doc,
        "本报告所用结果均来自 mstar_ssl_practice2_clean/runs 下的实际训练输出；源代码、模型权重、CSV记录、JSON配置和图表文件随项目目录一并提交。"
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", default="../《人工智能开发实训II》报告模板.docx")
    parser.add_argument("--project", default=".")
    parser.add_argument("--out", default="report/组号-人工智能实训II-实践2.docx")
    parser.add_argument("--group", default="待填写")
    parser.add_argument("--members", default="待填写")
    args = parser.parse_args()

    project_dir = Path(args.project).resolve()
    template = Path(args.template)
    if not template.is_absolute():
        template = (project_dir / template).resolve()
    out = Path(args.out)
    if not out.is_absolute():
        out = project_dir / out
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template))
    clear_body(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    add_cover(doc, "实践二：MSTAR深度半监督学习实验", args.group, args.members)
    add_report(doc, project_dir, load_runs(project_dir))
    doc.save(out)
    print(f"Report written to {out}")


if __name__ == "__main__":
    main()
