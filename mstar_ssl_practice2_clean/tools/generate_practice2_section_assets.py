#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
import numpy as np
import pandas as pd


RUNS = {
    "full": ("01_supervised_full", "Full supervised"),
    "limited": ("02_supervised_10percent", "10% supervised"),
    "fixmatch": ("03_fixmatch_10percent", "FixMatch"),
}


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def pp(value: float) -> str:
    return f"{value * 100:.2f}"


def load_runs(runs_dir: Path) -> dict:
    loaded = {}
    for key, (folder, label) in RUNS.items():
        run_dir = runs_dir / folder
        loaded[key] = {
            "label": label,
            "dir": run_dir,
            "args": load_json(run_dir / "args.json"),
            "metrics": load_json(run_dir / "metrics.json"),
            "history": pd.read_csv(run_dir / "history.csv"),
            "confusion": pd.read_csv(run_dir / "confusion_matrix.csv", index_col=0),
        }
    return loaded


def train_acc_column(history: pd.DataFrame) -> str:
    return "train_acc" if "train_acc" in history.columns else "labeled_train_acc"


def add_values(ax, values: list[float]) -> None:
    for i, value in enumerate(values):
        ax.text(i, value + 1.0, f"{value:.2f}%", ha="center", va="bottom", fontsize=9)


def save_accuracy_bar(runs: dict, keys: list[str], out: Path, title: str) -> None:
    labels = [runs[key]["label"] for key in keys]
    values = [runs[key]["metrics"]["best_test_acc"] * 100 for key in keys]
    colors = ["#4C78A8", "#F58518", "#54A24B"][: len(keys)]
    fig, ax = plt.subplots(figsize=(6.4, 3.6), dpi=220)
    ax.bar(labels, values, color=colors, width=0.55)
    ax.set_ylim(0, 105)
    ax.set_ylabel("Best test accuracy (%)")
    ax.set_title(title)
    ax.grid(axis="y", linestyle="--", alpha=0.35)
    add_values(ax, values)
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_curves(runs: dict, keys: list[str], out: Path, title: str) -> None:
    fig, axes = plt.subplots(2, 2, figsize=(9.4, 6.0), dpi=220)
    for key in keys:
        run = runs[key]
        hist = run["history"]
        label = run["label"]
        axes[0, 0].plot(hist["epoch"], hist["train_loss"], label=label)
        axes[0, 1].plot(hist["epoch"], hist["test_loss"], label=label)
        axes[1, 0].plot(hist["epoch"], hist[train_acc_column(hist)] * 100, label=label)
        axes[1, 1].plot(hist["epoch"], hist["test_acc"] * 100, label=label)
    for ax, subtitle in zip(
        axes.ravel(),
        ["Train loss", "Test loss", "Train accuracy", "Test accuracy"],
    ):
        ax.set_title(subtitle)
        ax.set_xlabel("Epoch")
        ax.grid(True, linestyle="--", alpha=0.3)
        ax.legend(fontsize=8)
    axes[1, 0].set_ylabel("%")
    axes[1, 1].set_ylabel("%")
    fig.suptitle(title, y=1.02, fontsize=12)
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_confusion(runs: dict, keys: list[str], out: Path, title: str) -> None:
    fig = plt.figure(figsize=(4.9 * len(keys) + 0.55, 4.6), dpi=220)
    grid = fig.add_gridspec(
        1,
        len(keys) + 1,
        width_ratios=[1] * len(keys) + [0.055],
        wspace=0.35,
    )
    axes = [fig.add_subplot(grid[0, i]) for i in range(len(keys))]
    cax = fig.add_subplot(grid[0, len(keys)])
    im = None
    for ax, key in zip(axes, keys):
        matrix = runs[key]["confusion"].to_numpy(dtype=float)
        row_sum = matrix.sum(axis=1, keepdims=True)
        norm = np.divide(matrix, row_sum, out=np.zeros_like(matrix), where=row_sum != 0)
        im = ax.imshow(norm, cmap="Blues", vmin=0, vmax=1)
        labels = list(runs[key]["confusion"].index)
        ax.set_title(runs[key]["label"])
        ax.set_xlabel("Predicted")
        ax.set_ylabel("True")
        ax.set_xticks(range(len(labels)), labels, rotation=90, fontsize=6)
        ax.set_yticks(range(len(labels)), labels, fontsize=6)
    if im is not None:
        fig.colorbar(im, cax=cax)
    fig.suptitle(title, y=0.98, fontsize=12)
    fig.subplots_adjust(top=0.82, bottom=0.24, left=0.06, right=0.95)
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_recovery(runs: dict, out: Path) -> None:
    values = [runs[key]["metrics"]["best_test_acc"] * 100 for key in ["full", "limited", "fixmatch"]]
    full, limited, fixmatch = values
    drop = full - limited
    gain = fixmatch - limited
    recovery = gain / drop * 100 if drop > 0 else 0.0

    fig, axes = plt.subplots(1, 2, figsize=(9.4, 3.6), dpi=220)
    axes[0].bar(["Full", "10% Sup.", "FixMatch"], values, color=["#4C78A8", "#F58518", "#54A24B"], width=0.55)
    axes[0].set_ylim(0, 105)
    axes[0].set_ylabel("Best test accuracy (%)")
    axes[0].set_title("Accuracy after adding SSL")
    axes[0].grid(axis="y", linestyle="--", alpha=0.35)
    add_values(axes[0], values)

    axes[1].plot(["Full", "10% Sup.", "FixMatch"], values, marker="o", color="#333333", linewidth=2)
    axes[1].set_ylim(0, 105)
    axes[1].set_ylabel("Best test accuracy (%)")
    axes[1].set_title("Accuracy loss and recovery")
    axes[1].grid(True, linestyle="--", alpha=0.35)
    axes[1].annotate(
        f"drop {drop:.2f} pp",
        xy=(0.5, (full + limited) / 2),
        xytext=(0.05, limited + 10),
        arrowprops={"arrowstyle": "->", "lw": 1},
        fontsize=8,
    )
    axes[1].annotate(
        f"gain {gain:.2f} pp\nrecovery {recovery:.1f}%",
        xy=(1.5, (limited + fixmatch) / 2),
        xytext=(1.35, fixmatch + 9),
        arrowprops={"arrowstyle": "->", "lw": 1},
        fontsize=8,
    )
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_fixmatch_flow(out: Path) -> None:
    font_path = Path("assets/fonts/NotoSansCJKsc-Regular.otf")
    font = FontProperties(fname=str(font_path)) if font_path.exists() else None
    fig, ax = plt.subplots(figsize=(9.4, 3.2), dpi=220)
    ax.axis("off")
    boxes = [
        (0.03, 0.62, "有标签图像\n弱增强"),
        (0.28, 0.62, "交叉熵\n监督损失"),
        (0.03, 0.18, "无标签图像\n弱增强 + 强增强"),
        (0.28, 0.18, "弱增强预测\n生成伪标签"),
        (0.52, 0.18, "置信度筛选\n最大概率 >= 0.95"),
        (0.75, 0.18, "强增强预测\n一致性损失"),
    ]
    for x, y, text in boxes:
        rect = plt.Rectangle((x, y), 0.18, 0.2, facecolor="#F7F7F7", edgecolor="#333333", linewidth=1)
        ax.add_patch(rect)
        ax.text(x + 0.09, y + 0.1, text, ha="center", va="center", fontsize=8, fontproperties=font)
    arrows = [
        ((0.21, 0.72), (0.28, 0.72)),
        ((0.21, 0.28), (0.28, 0.28)),
        ((0.46, 0.28), (0.52, 0.28)),
        ((0.70, 0.28), (0.75, 0.28)),
        ((0.46, 0.72), (0.75, 0.38)),
    ]
    for start, end in arrows:
        ax.annotate("", xy=end, xytext=start, arrowprops={"arrowstyle": "->", "lw": 1.2})
    ax.text(
        0.84,
        0.64,
        "总损失：\nL = Lx + λu Lu",
        ha="center",
        va="center",
        fontsize=10,
        fontproperties=font,
    )
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def markdown(runs: dict, out_dir: Path) -> str:
    full = runs["full"]["metrics"]["best_test_acc"]
    limited = runs["limited"]["metrics"]["best_test_acc"]
    fixmatch = runs["fixmatch"]["metrics"]["best_test_acc"]
    drop = full - limited
    gain = fixmatch - limited
    recovery = gain / drop if drop > 0 else 0.0
    content = f"""# 实践二：MSTAR 深度半监督学习实验

## 1. 基线 vs 有限标签

本实验首先训练全标签监督模型作为基线，再将训练标签减少到 10% 后训练同一网络。MSTAR 训练集共 2746 张图像，10% 标签设置下按类别分层保留 268 张有标签样本，其余 2478 张在纯监督训练中不使用。

| 实验 | 有标签样本 | 未使用无标签样本 | batch size | 学习率 | 最优测试精度 | 最优 epoch |
| --- | ---: | ---: | ---: | ---: | ---: | ---: |
| 全标签监督 | 2746 | 0 | 64 | 0.001 | {pct(full)} | {runs["full"]["metrics"]["best_epoch"]} |
| 10% 标签监督 | 268 | 2478 | 64 | 0.001 | {pct(limited)} | {runs["limited"]["metrics"]["best_epoch"]} |

![基线与有限标签精度对比](../report_assets_practice2/01_baseline_limited_accuracy_bar.png)

![基线与有限标签训练/测试曲线](../report_assets_practice2/02_baseline_limited_curves.png)

![基线与有限标签混淆矩阵](../report_assets_practice2/03_baseline_limited_confusion.png)

结果显示，全标签监督最优测试精度为 {pct(full)}，10% 标签监督下降到 {pct(limited)}，下降 {pp(drop)} 个百分点。训练曲线和混淆矩阵说明，标签减少后模型更容易受少量标注样本限制，测试集泛化能力明显下降。

## 2. 半监督方法

本实验采用 FixMatch。其核心思想是：对无标签样本生成弱增强和强增强两个视图，先用弱增强视图得到预测概率；当最大类别概率不低于阈值 0.95 时，将该预测类别作为伪标签，再约束强增强视图输出同一类别。训练损失为 `L = Lx + lambda_u Lu`，其中 `Lx` 是有标签交叉熵损失，`Lu` 是无标签一致性损失，本实验 `lambda_u=1.0`。

![FixMatch 实现框图](../report_assets_practice2/04_fixmatch_flowchart.png)

## 3. 加入半监督后的结果对比

FixMatch 与 10% 标签监督使用完全相同的 268 张有标签样本，额外把剩余 2478 张训练图像作为无标签样本。FixMatch 的有标签 batch size 为 64，`μ=2`，因此每个 step 同时使用 64 张有标签图像和 128 张无标签图像。

| 实验 | 有标签样本 | 无标签/未使用样本 | batch size | μ | 学习率 | 最优测试精度 | 最优 epoch |
| --- | ---: | ---: | ---: | ---: | ---: | ---: | ---: |
| 全标签监督 | 2746 | 0 | 64 | - | 0.001 | {pct(full)} | {runs["full"]["metrics"]["best_epoch"]} |
| 10% 标签监督 | 268 | 2478 | 64 | - | 0.001 | {pct(limited)} | {runs["limited"]["metrics"]["best_epoch"]} |
| FixMatch 半监督 | 268 | 2478 | 64 | 2 | 0.001 | {pct(fixmatch)} | {runs["fixmatch"]["metrics"]["best_epoch"]} |

![半监督精度恢复情况](../report_assets_practice2/05_ssl_accuracy_recovery.png)

![三组实验训练/测试曲线](../report_assets_practice2/06_all_training_curves.png)

![三组实验混淆矩阵](../report_assets_practice2/07_all_confusion_matrices.png)

加入 FixMatch 后，最优测试精度从 10% 标签监督的 {pct(limited)} 提升到 {pct(fixmatch)}，提高 {pp(gain)} 个百分点，恢复了约 {recovery * 100:.2f}% 的精度缺口。该结果说明无标签样本能够通过伪标签和一致性约束提供额外监督信号，从而缓解有限标签导致的性能下降。
"""
    path = out_dir / "practice2_section_2to3_pages.md"
    path.write_text(content, encoding="utf-8")
    return content


def html(runs: dict, out_dir: Path) -> None:
    full = runs["full"]["metrics"]["best_test_acc"]
    limited = runs["limited"]["metrics"]["best_test_acc"]
    fixmatch = runs["fixmatch"]["metrics"]["best_test_acc"]
    drop = full - limited
    gain = fixmatch - limited
    recovery = gain / drop if drop > 0 else 0.0
    content = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>实践二：MSTAR 深度半监督学习实验</title>
  <style>
    @page {{ size: A4; margin: 16mm 15mm; }}
    body {{
      font-family: "Microsoft YaHei", "SimSun", Arial, sans-serif;
      color: #111;
      font-size: 10.5pt;
      line-height: 1.45;
      max-width: 920px;
      margin: 0 auto;
    }}
    h1 {{ font-size: 18pt; margin: 0 0 8px; text-align: center; }}
    h2 {{ font-size: 13pt; margin: 12px 0 6px; }}
    p {{ margin: 5px 0; text-align: justify; }}
    table {{ width: 100%; border-collapse: collapse; margin: 6px 0 8px; font-size: 9pt; }}
    th, td {{ border: 1px solid #777; padding: 4px 5px; text-align: center; }}
    th {{ background: #f1f3f5; }}
    figure {{ margin: 7px 0 10px; page-break-inside: avoid; }}
    figcaption {{ text-align: center; font-size: 9pt; margin-top: 3px; color: #333; }}
    img {{ display: block; max-width: 100%; margin: 0 auto; }}
    .wide img {{ max-height: 102mm; }}
    .short img {{ max-height: 62mm; }}
    .page-break {{ page-break-before: always; }}
  </style>
</head>
<body>
  <h1>实践二：MSTAR 深度半监督学习实验</h1>

  <h2>1. 基线 vs 有限标签</h2>
  <p>本实验首先训练全标签监督模型作为基线，再将训练标签减少到 10% 后训练同一网络。MSTAR 训练集共 2746 张图像，10% 标签设置下按类别分层保留 268 张有标签样本，其余 2478 张在纯监督训练中不使用。</p>
  <table>
    <tr><th>实验</th><th>有标签样本</th><th>未使用无标签样本</th><th>batch size</th><th>学习率</th><th>最优测试精度</th><th>最优 epoch</th></tr>
    <tr><td>全标签监督</td><td>2746</td><td>0</td><td>64</td><td>0.001</td><td>{pct(full)}</td><td>{runs["full"]["metrics"]["best_epoch"]}</td></tr>
    <tr><td>10% 标签监督</td><td>268</td><td>2478</td><td>64</td><td>0.001</td><td>{pct(limited)}</td><td>{runs["limited"]["metrics"]["best_epoch"]}</td></tr>
  </table>
  <figure class="short"><img src="01_baseline_limited_accuracy_bar.png"><figcaption>图2-1 基线与有限标签精度对比</figcaption></figure>
  <figure class="wide"><img src="02_baseline_limited_curves.png"><figcaption>图2-2 基线与有限标签训练/测试曲线</figcaption></figure>
  <figure class="short"><img src="03_baseline_limited_confusion.png"><figcaption>图2-3 基线与有限标签混淆矩阵</figcaption></figure>
  <p>结果显示，全标签监督最优测试精度为 {pct(full)}，10% 标签监督下降到 {pct(limited)}，下降 {pp(drop)} 个百分点。训练曲线和混淆矩阵说明，标签减少后模型更容易受少量标注样本限制，测试集泛化能力明显下降。</p>

  <h2>2. 半监督方法</h2>
  <p>本实验采用 FixMatch。其核心思想是：对无标签样本生成弱增强和强增强两个视图，先用弱增强视图得到预测概率；当最大类别概率不低于阈值 0.95 时，将该预测类别作为伪标签，再约束强增强视图输出同一类别。训练损失为 L = Lx + lambda_u Lu，其中 Lx 是有标签交叉熵损失，Lu 是无标签一致性损失，本实验 lambda_u=1.0。</p>
  <figure class="short"><img src="04_fixmatch_flowchart.png"><figcaption>图2-4 FixMatch 实现框图</figcaption></figure>

  <h2>3. 加入半监督后的结果对比</h2>
  <p>FixMatch 与 10% 标签监督使用完全相同的 268 张有标签样本，额外把剩余 2478 张训练图像作为无标签样本。FixMatch 的有标签 batch size 为 64，μ=2，因此每个 step 同时使用 64 张有标签图像和 128 张无标签图像。</p>
  <table>
    <tr><th>实验</th><th>有标签样本</th><th>无标签/未使用样本</th><th>batch size</th><th>μ</th><th>学习率</th><th>最优测试精度</th><th>最优 epoch</th></tr>
    <tr><td>全标签监督</td><td>2746</td><td>0</td><td>64</td><td>-</td><td>0.001</td><td>{pct(full)}</td><td>{runs["full"]["metrics"]["best_epoch"]}</td></tr>
    <tr><td>10% 标签监督</td><td>268</td><td>2478</td><td>64</td><td>-</td><td>0.001</td><td>{pct(limited)}</td><td>{runs["limited"]["metrics"]["best_epoch"]}</td></tr>
    <tr><td>FixMatch 半监督</td><td>268</td><td>2478</td><td>64</td><td>2</td><td>0.001</td><td>{pct(fixmatch)}</td><td>{runs["fixmatch"]["metrics"]["best_epoch"]}</td></tr>
  </table>
  <figure class="short"><img src="05_ssl_accuracy_recovery.png"><figcaption>图2-5 半监督精度恢复情况</figcaption></figure>
  <figure class="wide"><img src="06_all_training_curves.png"><figcaption>图2-6 三组实验训练/测试曲线</figcaption></figure>
  <figure class="short"><img src="07_all_confusion_matrices.png"><figcaption>图2-7 三组实验混淆矩阵</figcaption></figure>
  <p>加入 FixMatch 后，最优测试精度从 10% 标签监督的 {pct(limited)} 提升到 {pct(fixmatch)}，提高 {pp(gain)} 个百分点，恢复了约 {recovery * 100:.2f}% 的精度缺口。该结果说明无标签样本能够通过伪标签和一致性约束提供额外监督信号，从而缓解有限标签导致的性能下降。</p>
</body>
</html>
"""
    (out_dir / "practice2_section_2to3_pages.html").write_text(content, encoding="utf-8")


def try_docx(markdown_text: str, out_dir: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Cm
    except Exception:
        return

    doc = Document()
    doc.add_heading("实践二：MSTAR 深度半监督学习实验", level=1)
    doc.add_heading("1. 基线 vs 有限标签", level=2)
    doc.add_paragraph(
        "全标签监督最优测试精度为 97.07%，10% 标签监督下降到 73.65%，下降 23.42 个百分点。"
    )
    doc.add_picture(str(out_dir / "01_baseline_limited_accuracy_bar.png"), width=Cm(13.5))
    doc.add_picture(str(out_dir / "02_baseline_limited_curves.png"), width=Cm(15.0))
    doc.add_picture(str(out_dir / "03_baseline_limited_confusion.png"), width=Cm(15.0))
    doc.add_heading("2. 半监督方法", level=2)
    doc.add_paragraph(
        "FixMatch 使用弱增强预测生成高置信伪标签，再约束强增强视图保持一致。总损失为 L = Lx + lambda_u Lu。"
    )
    doc.add_picture(str(out_dir / "04_fixmatch_flowchart.png"), width=Cm(15.0))
    doc.add_heading("3. 加入半监督后的结果对比", level=2)
    doc.add_paragraph(
        "FixMatch 与 10% 标签监督使用相同 268 张有标签样本，并额外利用 2478 张无标签图像。"
        "最优测试精度提升到 82.60%，相对 10% 标签监督提高 8.95 个百分点。"
    )
    doc.add_picture(str(out_dir / "05_ssl_accuracy_recovery.png"), width=Cm(15.0))
    doc.add_picture(str(out_dir / "06_all_training_curves.png"), width=Cm(15.0))
    doc.add_picture(str(out_dir / "07_all_confusion_matrices.png"), width=Cm(15.0))
    doc.add_paragraph("完整 Markdown 正文见 practice2_section_2to3_pages.md。")
    doc.save(out_dir / "practice2_section_2to3_pages.docx")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--runs", default="runs")
    parser.add_argument("--out", default="report_assets_practice2")
    args = parser.parse_args()

    runs = load_runs(Path(args.runs))
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    save_accuracy_bar(
        runs,
        ["full", "limited"],
        out_dir / "01_baseline_limited_accuracy_bar.png",
        "Baseline vs limited labels",
    )
    save_curves(
        runs,
        ["full", "limited"],
        out_dir / "02_baseline_limited_curves.png",
        "Baseline vs limited labels: train/test curves",
    )
    save_confusion(
        runs,
        ["full", "limited"],
        out_dir / "03_baseline_limited_confusion.png",
        "Baseline vs limited labels: confusion matrices",
    )
    save_fixmatch_flow(out_dir / "04_fixmatch_flowchart.png")
    save_recovery(runs, out_dir / "05_ssl_accuracy_recovery.png")
    save_curves(
        runs,
        ["full", "limited", "fixmatch"],
        out_dir / "06_all_training_curves.png",
        "Full supervised vs 10% supervised vs FixMatch",
    )
    save_confusion(
        runs,
        ["full", "limited", "fixmatch"],
        out_dir / "07_all_confusion_matrices.png",
        "Full supervised vs 10% supervised vs FixMatch",
    )
    text = markdown(runs, out_dir)
    html(runs, out_dir)
    try_docx(text, out_dir)
    print(f"Practice 2 section assets written to {out_dir.resolve()}")


if __name__ == "__main__":
    main()
